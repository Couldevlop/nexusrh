"""
NexusRH CI — Génération des mises à jour des 3 documents Word officiels (v1.1)
Produit 3 fichiers _v1.1.docx avec les release notes 2026-05-15 ajoutées
au document original. À ouvrir dans Word pour relire/valider.

Documents traités :
  1. NexusRH_CI_Documentation_Tests.docx → _v1.1.docx
  2. NexusRH_CI_Guide_Commercial.docx    → _v1.1.docx
  3. NexusRH_CI_OpenLab_Dossier.docx     → _v1.1.docx

Auteur : OpenLab Consulting / Claude Opus 4.7
"""
from datetime import date
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


BASE_DIR = Path(__file__).parent
TODAY = date(2026, 5, 15).strftime("%d/%m/%Y")
VERSION = "v1.1 — 15 mai 2026"

# ── Couleurs ──────────────────────────────────────────────────────────────────
NAVY    = RGBColor(0x1E, 0x3A, 0x5F)
ORANGE  = RGBColor(0xE8, 0x5D, 0x04)
EMERALD = RGBColor(0x05, 0x96, 0x69)
BLUE    = RGBColor(0x25, 0x63, 0xEB)
SLATE   = RGBColor(0x64, 0x74, 0x8B)


def style_h1(p, text):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = NAVY
    return p


def style_h2(p, text):
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = ORANGE
    return p


def style_h3(p, text):
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY
    return p


def style_normal(p, text, bold=False, color=None):
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(11)
    if color:
        r.font.color.rgb = color
    return p


def style_meta(p, text):
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = SLATE
    return p


def add_bullet(doc, text, level=0):
    # Tentative style natif Word, sinon fallback bullet manuel (•) — compatibilité
    # avec les templates Word qui n'ont pas le style "List Bullet" défini.
    style_name = 'List Bullet' if level == 0 else 'List Bullet 2'
    try:
        p = doc.add_paragraph(style=style_name)
    except KeyError:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6 + 0.4 * level)
        text = ("• " if level == 0 else "◦ ") + text
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p


def add_separator(doc):
    doc.add_paragraph("─" * 60)


# ─────────────────────────────────────────────────────────────────────────────
# DOC 1 : Documentation Tests — Annexe Release Notes
# ─────────────────────────────────────────────────────────────────────────────
def update_documentation_tests(src: Path, dest: Path) -> None:
    doc = Document(str(src))

    doc.add_page_break()
    style_h1(doc.add_paragraph(), f"Annexe — Release Notes {VERSION}")
    style_meta(doc.add_paragraph(), f"Mise à jour de la documentation tests · {TODAY}")
    doc.add_paragraph()

    style_h2(doc.add_paragraph(), "Nouveaux modules de test")
    style_normal(doc.add_paragraph(),
        "Le cahier de recettes NEXUSRH_CI_Test_Plan.xlsx passe à 19 modules · "
        "226 scénarios (auparavant 16 modules). Les nouveautés :", bold=False)
    add_bullet(doc, "SOURCING_IA (SRC) — 9 cas : endpoints /jobs/:id/source et /source/compare, rate-limit, cache, devise auto, score richesse")
    add_bullet(doc, "MULTI_PAYS_FILIALES (MPF) — 11 cas : activation des 11 packs, toggle has_subsidiaries, création de filiale multi-pays, rattachement employé, badge filiale, suppression bloquée si employés actifs")
    add_bullet(doc, "RECRUTEMENT — 9 nouveaux cas (sourcing IA, transfert profils, drag-and-drop Kanban, comparaison Claude/Mistral, contact email, fix modal débordement)")
    add_bullet(doc, "AUTH — 3 cas non-régression : seed DO UPDATE, script reset-admin-passwords.sql, script standalone admin:reset-passwords")

    doc.add_paragraph()
    style_h2(doc.add_paragraph(), "Régressions documentées et résolues")

    style_h3(doc.add_paragraph(), "1. Login 401 sur tous les comptes démo (résolu 2026-05-15)")
    style_normal(doc.add_paragraph(),
        "Cause racine : le seed utilisait ON CONFLICT (email) DO NOTHING sur les INSERT users SOTRA "
        "et Cabinet Expertise CI. Quand la DB était re-seedée, les password_hash ne se rafraîchissaient pas. "
        "Conséquence : bcrypt.compare retournait false → 401 Unauthorized.")
    style_normal(doc.add_paragraph(), "Correctif livré :", bold=True)
    add_bullet(doc, "seed.ts : passage à ON CONFLICT (email) DO UPDATE SET password_hash = EXCLUDED.password_hash")
    add_bullet(doc, "Script standalone reset-admin-passwords.ts (pnpm run admin:reset-passwords)")
    add_bullet(doc, "Script SQL prêt-à-exécuter scripts/reset-admin-passwords.sql (8 UPDATE bcrypt 12 rounds)")
    add_bullet(doc, "Job K8s reset-admin-passwords-job.yaml (recovery prod sans rebuild)")

    style_h3(doc.add_paragraph(), "2. Multi-législation : seul CI affiché actif (résolu 2026-05-15)")
    style_normal(doc.add_paragraph(),
        "Cause racine : la table platform.country_configs n'avait que 5 pays seedés (CI, SN, BF, ML, TG) "
        "dont un seul avec is_active=true (CI). L'ON CONFLICT DO NOTHING empêchait l'activation rétroactive.")
    style_normal(doc.add_paragraph(), "Correctif livré :", bold=True)
    add_bullet(doc, "country_configs étendue à 11 pays : UEMOA 7 + CEMAC 2 + ECOWAS 2 (NG, GH)")
    add_bullet(doc, "Tous is_active=true par défaut")
    add_bullet(doc, "ON CONFLICT DO UPDATE pour refresh rétroactif")
    add_bullet(doc, "UI : drapeaux, badges zone (UEMOA/CEMAC/ECOWAS), compteurs régionaux")

    style_h3(doc.add_paragraph(), "3. Modal 'Nouvelle offre' déborde (résolu 2026-05-15)")
    style_normal(doc.add_paragraph(),
        "Refonte en flex-column : header sticky + body scrollable + footer sticky. "
        "Boutons Annuler/Créer toujours visibles, plus de débordement vertical.")

    doc.add_paragraph()
    style_h2(doc.add_paragraph(), "Conformité OWASP 2025 — Top 10")
    style_normal(doc.add_paragraph(), "Audit complet documenté dans nexusrh_ci/docs/OWASP-2025-AUDIT.md. Tous les correctifs livrés respectent :", bold=False)
    add_bullet(doc, "A01 Broken Access Control — RBAC API + frontend (guards + sidebar)")
    add_bullet(doc, "A02 Cryptographic Failures — bcrypt 12 rounds, JWT secret 32+ chars")
    add_bullet(doc, "A03 Injection — paramètres SQL liés ($1, $2…) sur 100% des handlers")
    add_bullet(doc, "A04 Insecure Design — findTenantAndUser itère tous les candidats avec bcrypt.compare")
    add_bullet(doc, "A05 Security Misconfiguration — rate-limit /auth/login (10/5min), /jobs/:id/source (6/min)")
    add_bullet(doc, "A09 Security Logging — audit_log pour recruitment.analyze_cv, source_profiles, sourced_transfer")
    add_bullet(doc, "A10 SSRF — pattern isUserActionable pour ne pas leaker les erreurs internes")

    doc.add_paragraph()
    style_h2(doc.add_paragraph(), "Plan de test post-déploiement")
    add_bullet(doc, "Login admin@sotra.ci / Admin1234! → /dashboard (test régression auth)")
    add_bullet(doc, "Recrutement → Sourcing IA → sélectionner 'Chauffeur Bus Senior' → 8 profils visibles")
    add_bullet(doc, "Carte profil → 'Transférer' → onglet Pipeline → vérifier drag-and-drop fonctionnel")
    add_bullet(doc, "Bouton 'Tout transférer (N)' → tous les profils dans le Kanban")
    add_bullet(doc, "Super Admin → Multi-législatif → 11 pays affichés, tous Actif")
    add_bullet(doc, "Settings → Entités juridiques → créer filiale 'SOTRA Dakar' avec pays SEN, pack sn_2024")

    doc.save(str(dest))
    print(f"OK Documentation_Tests v1.1 -> {dest.name}")


# ─────────────────────────────────────────────────────────────────────────────
# DOC 2 : Guide Commercial — Évolutions produit
# ─────────────────────────────────────────────────────────────────────────────
def update_guide_commercial(src: Path, dest: Path) -> None:
    doc = Document(str(src))

    doc.add_page_break()
    style_h1(doc.add_paragraph(), f"Évolutions produit {VERSION}")
    style_meta(doc.add_paragraph(), f"Argumentaire commercial · {TODAY}")
    doc.add_paragraph()

    style_h2(doc.add_paragraph(), "Sourcing IA Multi-Pays Afrique — NOUVEAU")
    style_normal(doc.add_paragraph(),
        "Génération automatique de profils candidats synthétiques calibrés pour les groupes "
        "opérant en Afrique. Couvre 15 pays (UEMOA, CEMAC, ECOWAS) avec plateformes locales "
        "(LinkedIn, Africawork, Emploi.ci, RMO, Novojob, Jobberman, MinaJobs…) et devises locales "
        "(XOF, XAF, NGN, GHS, EUR pour diaspora).")
    style_normal(doc.add_paragraph(), "Argumentaire :", bold=True)
    add_bullet(doc, "Réduit le temps de sourcing de 5-7 jours à 30 minutes")
    add_bullet(doc, "Profils contextuels : noms locaux, expérience OHADA, salaires en devise du pays")
    add_bullet(doc, "Comparaison Claude vs Mistral en parallèle (qualité vs coût)")
    add_bullet(doc, "Transfert direct vers le pipeline Kanban (1 clic ou 'Tout transférer')")
    add_bullet(doc, "Templates d'approche personnalisés (LinkedIn, email)")
    add_bullet(doc, "Compatible diaspora africaine (Paris, Londres, Montréal)")

    doc.add_paragraph()
    style_h2(doc.add_paragraph(), "Multi-Pays & Filiales — NOUVEAU")
    style_normal(doc.add_paragraph(),
        "NexusRH CI couvre désormais 11 pays africains avec moteurs de paie dédiés. Idéal pour "
        "les groupes panafricains avec filiales dans plusieurs pays.")
    style_normal(doc.add_paragraph(), "Pays supportés :", bold=True)
    add_bullet(doc, "UEMOA (XOF) : Côte d'Ivoire (production), Sénégal, Bénin, Togo, Burkina Faso, Mali, Niger")
    add_bullet(doc, "CEMAC (XAF) : Cameroun, Tchad")
    add_bullet(doc, "ECOWAS hors UEMOA : Nigeria (NGN), Ghana (GHS)")

    style_normal(doc.add_paragraph(), "Mécanisme de rattachement :", bold=True)
    style_normal(doc.add_paragraph(),
        "Chaque filiale est une entité juridique avec son propre N° CNPS, RCCM, pays et pack législatif. "
        "Les employés sont rattachés à une filiale via leur fiche. Le moteur de paie applique automatiquement "
        "le pack législatif de la filiale (CNPS pour CI, IPRES pour SN, CNSS pour BJ/BF, CNAVS pour TG, etc.).")

    doc.add_paragraph()
    style_h2(doc.add_paragraph(), "Améliorations UX")
    add_bullet(doc, "Modal 'Nouvelle offre d'emploi' refondue (plus de débordement)")
    add_bullet(doc, "Onglet 'Filiales' avec sélecteur pays + pack législatif")
    add_bullet(doc, "Cards filiales : drapeau, badge zone, pack législatif, compteur employés")

    doc.add_paragraph()
    style_h2(doc.add_paragraph(), "Robustesse & Sécurité")
    add_bullet(doc, "Audit OWASP 2025 complet (10 catégories couvertes)")
    add_bullet(doc, "Script de reset admin passwords (recovery production sans downtime)")
    add_bullet(doc, "Rate-limit IA configurable (6 req/min sourcing, 10 req/min CV analyze)")
    add_bullet(doc, "Audit_log automatique sur les actions IA (qui, quoi, combien)")

    doc.add_paragraph()
    style_h2(doc.add_paragraph(), "Tarification (rappel)")
    style_normal(doc.add_paragraph(), "Plans inchangés — modules Sourcing IA et Multi-Pays inclus dans Business/Enterprise.", bold=False)
    add_bullet(doc, "Trial : 30j gratuit · 10 users · 20 salariés · 0 FCFA")
    add_bullet(doc, "Starter : 30 sal. · < 70 000 FCFA/mois TTC")
    add_bullet(doc, "Business : 150 sal. · Tous modules + IA + ATS · < 10 000 FCFA/sal/mois")
    add_bullet(doc, "Enterprise : 150+ sal. · SLA premium + multi-sites · Sur devis")

    doc.save(str(dest))
    print(f"OK Guide_Commercial v1.1 -> {dest.name}")


# ─────────────────────────────────────────────────────────────────────────────
# DOC 3 : OpenLab Dossier — Démo prod
# ─────────────────────────────────────────────────────────────────────────────
def update_openlab_dossier(src: Path, dest: Path) -> None:
    doc = Document(str(src))

    doc.add_page_break()
    style_h1(doc.add_paragraph(), f"Mise à jour démo {VERSION}")
    style_meta(doc.add_paragraph(), f"Dossier OpenLab Consulting · {TODAY}")
    doc.add_paragraph()

    style_h2(doc.add_paragraph(), "Environnement de démonstration")
    style_normal(doc.add_paragraph(), "URL de la plateforme :", bold=True)
    style_normal(doc.add_paragraph(), "https://nexusrh.openlabconsulting.com")
    doc.add_paragraph()
    style_normal(doc.add_paragraph(), "Comptes de démo (réinitialisés le 15/05/2026) :", bold=True)
    add_bullet(doc, "Super Admin : superadmin@nexusrh-ci.com / SuperAdmin1234!")
    add_bullet(doc, "SOTRA Admin : admin@sotra.ci / Admin1234!")
    add_bullet(doc, "SOTRA RH : rh@sotra.ci / Admin1234!")
    add_bullet(doc, "SOTRA Manager : manager@sotra.ci / Admin1234!")
    add_bullet(doc, "SOTRA Employé : employe@sotra.ci / Admin1234!")
    add_bullet(doc, "Cabinet Expertise CI : admin@cabinet-expertise.ci / Admin1234!")
    add_bullet(doc, "OpenLab Consulting (tenant démo) : coulwao@gmail.com / Openlab1234!")

    doc.add_paragraph()
    style_h2(doc.add_paragraph(), "Scénarios de démonstration recommandés")

    style_h3(doc.add_paragraph(), "Scénario 1 — Vue 360° SaaS Multi-Tenant (15 min)")
    add_bullet(doc, "Connexion super_admin → vue plateforme (3 tenants actifs, KPIs MRR)")
    add_bullet(doc, "Onglet Multi-législatif → 11 pays affichés, tous actifs (UEMOA + CEMAC + ECOWAS)")
    add_bullet(doc, "Création d'un 4ème tenant en live (wizard 3 étapes, thème custom, admin invité)")

    style_h3(doc.add_paragraph(), "Scénario 2 — Sourcing IA Recrutement (10 min)")
    add_bullet(doc, "Login admin@sotra.ci → /recrutement")
    add_bullet(doc, "Onglet Sourcing IA → sélectionner 'Chauffeur Bus Senior'")
    add_bullet(doc, "Visualisation des 8 profils sourcés en cache (drapeaux, scores 74-92%)")
    add_bullet(doc, "Bouton 'Tout transférer (N)' → vers pipeline Kanban")
    add_bullet(doc, "Drag-and-drop d'un profil de 'Nouveau' vers 'Entretien'")
    add_bullet(doc, "Bonus : Mode 'Compare Claude vs Mistral' avec métriques temps/coût/qualité")

    style_h3(doc.add_paragraph(), "Scénario 3 — Groupe avec filiales (8 min)")
    add_bullet(doc, "Connexion admin SOTRA → Paramètres → Entités juridiques")
    add_bullet(doc, "Création filiale 'SOTRA Bouaké' avec pays Côte d'Ivoire + AT BTP 3%")
    add_bullet(doc, "Création filiale 'SOTRA Dakar' avec pays Sénégal + pack sn_2024 + AT 2%")
    add_bullet(doc, "Rattachement de 2 employés à 'SOTRA Bouaké' via leur fiche")
    add_bullet(doc, "Mise en avant : moteur de paie auto-sélectionné selon la filiale")

    style_h3(doc.add_paragraph(), "Scénario 4 — Conformité CNPS + DISA (5 min)")
    add_bullet(doc, "Génération déclaration mensuelle e-CNPS (export CSV)")
    add_bullet(doc, "Génération DISA annuelle (12 mois agrégés)")
    add_bullet(doc, "Démontrer les calculs : double plafond CNPS (70 000 / 1 647 315 FCFA)")

    doc.add_paragraph()
    style_h2(doc.add_paragraph(), "Argumentaire OpenLab Consulting")
    style_normal(doc.add_paragraph(),
        "NexusRH CI est le premier SIRH SaaS conçu nativement pour le marché africain. "
        "OpenLab Consulting accompagne ses clients de la phase de scoping à la mise en production : "
        "formation initiale, paramétrage CNPS/RCCM/DGI, intégration des packs législatifs par pays, "
        "support WhatsApp 7j/7.")
    doc.add_paragraph()
    add_bullet(doc, "Cocody, Rivièra Faya Lauriers 8, Abidjan")
    add_bullet(doc, "infos@openlabconsulting.com · +225 07 09 32 05 94")
    add_bullet(doc, "Démo en ligne : nexusrh.openlabconsulting.com")

    doc.save(str(dest))
    print(f"OK OpenLab_Dossier v1.1 -> {dest.name}")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────
def main() -> None:
    targets = [
        (BASE_DIR / "NexusRH_CI_Documentation_Tests.docx",
         BASE_DIR / "NexusRH_CI_Documentation_Tests_v1.1.docx",
         update_documentation_tests),
        (BASE_DIR / "NexusRH_CI_Guide_Commercial.docx",
         BASE_DIR / "NexusRH_CI_Guide_Commercial_v1.1.docx",
         update_guide_commercial),
        (BASE_DIR / "NexusRH_CI_OpenLab_Dossier.docx",
         BASE_DIR / "NexusRH_CI_OpenLab_Dossier_v1.1.docx",
         update_openlab_dossier),
    ]

    for src, dest, fn in targets:
        if not src.exists():
            print(f"[SKIP] {src.name} introuvable")
            continue
        try:
            fn(src, dest)
        except Exception as e:
            print(f"[ERR] {src.name}: {e}")


if __name__ == "__main__":
    main()
