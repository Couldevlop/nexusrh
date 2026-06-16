# -*- coding: utf-8 -*-
"""
Generateur du canevas de l'offre technique (Word) - DAO VERSUS BANK / NexusRH CI.
Dependance : python-docx (installe).
Sortie : docs/DAO/Canevas-Offre-Technique-NexusRH-CI.docx

Structure conforme : grille de notation du DAO + Clean Architecture + OWASP Top 10.
Les sections [A RENSEIGNER] sont a completer par l'equipe OpenLab.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "Canevas-Offre-Technique-NexusRH-CI.docx")
LOGO = os.path.normpath(os.path.join(HERE, "..", "..", "..", "OPENLAB.PNG"))  # nexusrh/OPENLAB.PNG

NAVY = RGBColor(0x0F, 0x2A, 0x44)
ACCENT = RGBColor(0xE8, 0x5D, 0x04)
GRAY = RGBColor(0x55, 0x5B, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
NAVY_HEX = "0F2A44"
ACCENT_HEX = "E85D04"
LIGHT_HEX = "F2F4F7"
ALT_HEX = "FAFBFC"

FONT = "Calibri"


# ---------- helpers bas niveau ----------
def shade(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)

def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcpr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement("w:%s" % tag)
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcpr.append(m)

def run(par, text, size=10.5, bold=False, italic=False, color=None, font=FONT):
    r = par.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r


# ---------- blocs de contenu ----------
def h1(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    # bandeau : tableau 1x1 fond navy
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    shade(cell, NAVY_HEX)
    set_cell_margins(cell, top=80, bottom=80, left=140, right=120)
    cp = cell.paragraphs[0]
    run(cp, text, size=14, bold=True, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    run(p, text, size=12.5, bold=True, color=NAVY)
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run(p, text, size=11, bold=True, color=ACCENT)
    return p

def body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run(p, text, size=10.5, italic=italic, color=RGBColor(0x1A, 0x1A, 0x1A))
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run(p, bold_prefix, size=10.5, bold=True, color=NAVY)
        run(p, text, size=10.5, color=RGBColor(0x1A, 0x1A, 0x1A))
    else:
        run(p, text, size=10.5, color=RGBColor(0x1A, 0x1A, 0x1A))
    return p

def callout(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    shade(cell, LIGHT_HEX)
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run(cp, text, size=10, italic=True, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl

def table(doc, header, rows, widths_cm):
    tbl = doc.add_table(rows=1, cols=len(header))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    # header
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(header):
        shade(hdr[i], NAVY_HEX)
        set_cell_margins(hdr[i])
        cp = hdr[i].paragraphs[0]
        run(cp, htext, size=9.5, bold=True, color=WHITE)
    # rows
    for ridx, r in enumerate(rows):
        cells = tbl.add_row().cells
        for i, val in enumerate(r):
            if ridx % 2 == 1:
                shade(cells[i], ALT_HEX)
            set_cell_margins(cells[i])
            cp = cells[i].paragraphs[0]
            run(cp, val, size=9.5, color=RGBColor(0x1A, 0x1A, 0x1A))
    # widths
    for i, w in enumerate(widths_cm):
        for row in tbl.rows:
            row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


# ---------- document ----------
def build():
    doc = Document()
    # marges + police de base
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)

    # ---- En-tete de document (header) ----
    hp = sec.header.paragraphs[0]
    run(hp, "OpenLab Consulting  -  NexusRH CI", size=8.5, bold=True, color=NAVY)
    hp.add_run("\t\tCanevas de l'offre technique").font.size = Pt(8)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(fp, "DAO DRH/SIRH/29042026  -  VERSUS BANK  -  Confidentiel", size=8, color=GRAY)

    # ---- Page de garde ----
    doc.add_paragraph()
    if os.path.exists(LOGO):
        lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.add_run().add_picture(LOGO, width=Cm(5.5))
        doc.add_paragraph()
    for _ in range(2):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(t, "OFFRE TECHNIQUE", size=13, bold=True, color=ACCENT)
    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(t2, "Canevas de redaction - NexusRH CI", size=26, bold=True, color=NAVY)
    t3 = doc.add_paragraph(); t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(t3, "Trame conforme au DAO, a la Clean Architecture et aux recommandations OWASP",
        size=12, italic=True, color=ACCENT)
    doc.add_paragraph()
    meta = [
        ("Reference", "DAO DRH/SIRH/29042026 - Projet TRANSFORMATION RH"),
        ("Autorite contractante", "VERSUS BANK (Cote d'Ivoire)"),
        ("Soumissionnaire", "OpenLab Consulting - solution NexusRH CI"),
        ("Objet", "Structure type de l'offre technique a remettre (notation /100)"),
        ("Cadre", "Langue : francais  -  Monnaie : FCFA (XOF)  -  Delai : 9 mois"),
    ]
    mt = doc.add_table(rows=0, cols=2)
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in meta:
        cells = mt.add_row().cells
        set_cell_margins(cells[0]); set_cell_margins(cells[1])
        run(cells[0].paragraphs[0], k, size=10, bold=True, color=NAVY)
        run(cells[1].paragraphs[0], v, size=10, color=GRAY)
        cells[0].width = Cm(5.0); cells[1].width = Cm(11.5)
    doc.add_page_break()

    # ---- Mode d'emploi ----
    h1(doc, "Mode d'emploi de ce canevas")
    body(doc, "Ce document est une trame de redaction. Chaque chapitre correspond a un critere de notation "
              "du DAO ou a une question posee aux soumissionnaires. Les sections [A RENSEIGNER] doivent etre "
              "completees par l'equipe OpenLab (chiffres, CV, plannings, prix). La structure suit l'ordre de la "
              "grille d'evaluation technique afin que la COJO retrouve facilement chaque point note.")
    table(doc, ["Chapitre du dossier technique", "Critere DAO", "Points"],
          [["Ch. 1 - Comprehension du besoin et perimetre", "Contexte", "-"],
           ["Ch. 2 - Couverture fonctionnelle (modules 1 a 9)", "Critere 1", "35"],
           ["Ch. 3 - Architecture, securite et conformite", "Critere 2", "20"],
           ["Ch. 4 - Interoperabilite et migration", "Critere 3", "15"],
           ["Ch. 5 - Methodologie et planning (9 mois)", "Critere 4", "10"],
           ["Ch. 6 - Conduite du changement, formation, support", "Critere 5", "10"],
           ["Ch. 7 - Equipe dediee et references", "Critere 6", "10"],
           ["Ch. 8 - Reponses aux questions du DAO (1 a 10)", "Annexe", "-"],
           ["Ch. 9 - Reversibilite, propriete des donnees, livrables", "CCAP", "-"],
           ["Annexe - Matrice de conformite fonctionnelle", "Annexe 2", "-"]],
          widths_cm=[10.5, 4.0, 2.0])

    # ---- Ch.1 ----
    doc.add_page_break()
    h1(doc, "Chapitre 1 - Comprehension du besoin et perimetre")
    h3(doc, "1.1  Rappel du contexte VERSUS BANK")
    body(doc, "Reformuler le contexte (modernisation, centralisation, securisation des RH ; suppression des "
              "traitements manuels Excel/papier ; environnement bancaire reglemente BCEAO ; tracabilite et "
              "conformite). Demontrer la comprehension du cycle de vie du collaborateur de bout en bout.")
    h3(doc, "1.2  Perimetre couvert et engagement de resultat")
    body(doc, "Lister les 13 domaines du perimetre (recrutement/integration, administration, paie, "
              "organigrammes, temps/absences, evaluations/objectifs, formation, carrieres/talents, mobilites, "
              "successions, climat social, reporting, interfaces). Rappeler les KPI cibles : 90% des dossiers "
              "numerises, 90% des demandes RH digitalisees, -50% de temps administratif a 6 mois, 90% "
              "d'utilisateurs actifs a 3 mois, -90% d'erreurs de saisie a 4 mois.")
    h3(doc, "1.3  Strategie de reponse : lot unique, deploiement par lots")
    body(doc, "Justifier le lot unique (tiers de confiance unique, coherence d'architecture, absence de rupture "
              "de flux) tout en proposant un deploiement par lots sur 9 mois (voir Ch. 5).")

    # ---- Ch.2 ----
    doc.add_page_break()
    h1(doc, "Chapitre 2 - Couverture fonctionnelle (35 points)")
    body(doc, "Pour chaque module, repondre selon le format DAO : (i) Oui/Non standard, (ii) description et "
              "limites, (iii) prerequis, (iv) elements de preuve (captures, doc), (v) impacts cout/delai/risque. "
              "Distinguer clairement le parametrage (standard) du developpement specifique (livre par lot).")
    for title, txt in [
        ("2.1  Recrutement, Onboarding et Offboarding",
         "Standard : ATS avec scoring IA, vivier, pipeline, page publique, pre-onboarding/checklists. "
         "Specifique (lot) : workflow de validation des besoins (Manager > DRH/Finance > DG), bibliotheque "
         "video et videos journalieres, bilan de fin d'integration, module offboarding + solde de tout compte."),
        ("2.2  Administration et dossier salarie",
         "Standard : dossier complet (etat civil, CNPS/NNI, contrats OHADA), historique des mouvements. "
         "Specifique : generateur multi-modeles de documents + signature electronique, gestion disciplinaire "
         "(niveau 4), suivi des visites medicales/aptitudes, coffre documentaire."),
        ("2.3  Temps, absences et conges (point fort)",
         "Standard : self-service web/mobile, typologie parametrable, calcul des soldes Code du travail CI "
         "(jours ouvrables, feries 2024), workflow N+1 puis RH, planning d'equipe."),
        ("2.4  Paie (deux options au choix de VERSUS BANK)",
         "NexusRH CI embarque un moteur de paie CI complet : laisser la paie a SAGE est OPTIONNEL. "
         "Option A - Paie native NexusRH : le moteur integre calcule la paie (CNPS double plafond, ITS/DGI, "
         "bulletins, declarations CNPS + DISA) ; SAGE devient facultatif. Option B - Interface SAGE "
         "(amont-paie) : SAGE garde le calcul, le SIRH collecte/controle/valide les variables et archive les "
         "bulletins. Specifique (lot 1, uniquement si Option B) : connecteur SAGE bidirectionnel "
         "(API/fichiers/SFTP) + coffre-fort des bulletins. Presenter les deux scenarios avec leurs impacts "
         "cout/delai/risque."),
        ("2.5  Talents, carrieres et competences",
         "Standard : matrice 9-box. Specifique : referentiel postes/competences (taxonomie de Bloom), "
         "campagnes d'evaluation + objectifs + validation N+2 + signature, calibrage, mobilites, successions/pools."),
        ("2.6  Formation",
         "Standard : catalogue, sessions, inscriptions, eligibilite FDFP. Specifique : plan de formation + "
         "workflow RH-DG + budget, evaluation a chaud/a froid, gestion des presences, interface e-learning (SSO)."),
        ("2.7  Organigrammes, reporting et portail",
         "Standard : reporting/dashboards (effectifs, turnover, absenteisme, masse salariale en FCFA), portail "
         "self-service mobile/PWA. Specifique : organigramme dynamique + export PDF/image, export generique "
         "Excel/CSV, enquetes climat social."),
    ]:
        h3(doc, title)
        body(doc, txt)
    callout(doc, "Joindre en annexe la Matrice de conformite fonctionnelle (Annexe 2 du DAO) entierement "
                 "renseignee, coherente avec les engagements de lot ci-dessus.")

    # ---- Ch.3 ----
    doc.add_page_break()
    h1(doc, "Chapitre 3 - Architecture, securite et conformite (20 points)")
    h2(doc, "3.1  Architecture applicative - principes Clean Architecture")
    body(doc, "Presenter une architecture en couches respectant la separation des responsabilites et la regle "
              "de dependance (les couches externes dependent des couches internes, jamais l'inverse) :")
    table(doc, ["Couche", "Responsabilite", "Mise en oeuvre NexusRH CI"],
          [["Domain (entites)", "Regles metier RH pures, independantes du framework",
            "Types et regles de paie CI (CNPS, ITS), conges, contrats OHADA - sans dependance Fastify/DB"],
           ["Use cases (services)", "Orchestration applicative, transactions",
            "Services : payroll-engine-ci, cnps, security-policy, recruitment-ai (logique isolee)"],
           ["Interface adapters", "Controleurs, presenters, repositories",
            "Routes Fastify (modules/*), validateurs Zod, repositories Drizzle"],
           ["Frameworks & drivers", "DB, web, IA, stockage, externes",
            "PostgreSQL/Drizzle, Redis, Anthropic SDK, S3/MinIO, Meilisearch, Mobile Money"]],
          widths_cm=[3.4, 5.0, 8.1])
    bullet(doc, "schema-per-tenant PostgreSQL, search_path par requete, nom de schema valide par liste blanche (anti-injection) ;", "Isolation multi-tenant : ")
    bullet(doc, "la logique metier (moteur de paie, regles) est testee unitairement sans infrastructure ; golden tests UI<->API ;", "Testabilite : ")
    bullet(doc, "la couche domaine ne connait ni Fastify ni Drizzle - elle est reutilisable (worker BullMQ, simulations) ;", "Frontiere claire : ")
    bullet(doc, "chaque module (employees, payroll, absences...) expose une API stable, modules activables par tenant (feature flags).", "Modularite : ")

    h2(doc, "3.2  Securite - conformite OWASP Top 10 (2021)")
    body(doc, "Demontrer, controle par controle, la couverture du Top 10 OWASP - exigence implicite d'un SI bancaire :")
    table(doc, ["Risque OWASP", "Mesure NexusRH CI", "Statut"],
          [["A01 Broken Access Control", "RBAC par role + perimetre manager fail-closed, isolation tenant, guards UI", "En place / a renforcer (role Admin Systeme)"],
           ["A02 Cryptographic Failures", "AES-256-GCM au repos (NNI/IBAN), TLS en transit, bcrypt 12 rounds", "En place"],
           ["A03 Injection", "Validation Zod stricte, requetes parametrees, liste blanche schema", "En place"],
           ["A04 Insecure Design", "Clean architecture, separation des taches (SoD) paie, revue de conception", "En place"],
           ["A05 Security Misconfiguration", "Secrets en variables d'env, durcissement headers, CORS controle", "En place"],
           ["A06 Vulnerable Components", "Audit dependances, mises a jour, SBOM a fournir", "A documenter"],
           ["A07 Auth Failures", "MFA TOTP + anti-rejeu, politique mdp, lockout, verif HIBP", "En place"],
           ["A08 Integrity Failures", "Piste d'audit non alterable, signature des livrables CI/CD", "En place / a etendre"],
           ["A09 Logging & Monitoring", "audit_log (qui/quoi/quand, avant/apres, IP) ; export SIEM a ajouter", "A completer (SIEM)"],
           ["A10 SSRF", "Garde SSRF sur connecteurs sortants (module integrations)", "En place"]],
          widths_cm=[4.2, 8.3, 4.0])

    h2(doc, "3.3  Habilitations - 6 profils RBAC et classification a 4 niveaux")
    body(doc, "Decrire le modele de droits : RBAC (roles) + perimetre (entite/agence) + relation manageriale "
              "(N+1/N+2) + regles exceptionnelles (interim, delegation). Aligner sur les 6 profils du DAO et "
              "engager le developpement du role Administrateur Systeme (acces technique, masquage des salaires) "
              "et de la classification native a 4 niveaux (Public / Interne / Confidentiel / Restreint) avec "
              "cloisonnement au sein du dossier salarie (identite vs remuneration vs sanctions vs sante) et "
              "controle des exports par niveau.")

    h2(doc, "3.4  Authentification, chiffrement, journalisation, PRA/PCA")
    bullet(doc, "SSO SAML/OIDC + Active Directory/Azure AD (a integrer), MFA, politique mdp, verrouillage, comptes inactifs ;", "Authentification : ")
    bullet(doc, "TLS en transit ; AES-256-GCM au repos ; gestion et rotation des cles (KMS) a formaliser ; sauvegardes chiffrees ;", "Chiffrement : ")
    bullet(doc, "contenu minimal des logs, protection et duree de conservation ; export syslog/API vers le SIEM de la banque (a developper) ;", "Journalisation et SIEM : ")
    bullet(doc, "RPO/RTO cibles, sauvegardes, tests de restauration, redondance, bascule (document a produire) ; dispo 99,5%, temps de reponse < 3 s.", "PRA/PCA : ")

    # ---- Ch.4 ----
    doc.add_page_break()
    h1(doc, "Chapitre 4 - Interoperabilite et migration (15 points)")
    h3(doc, "4.1  Interface SAGE (cible — uniquement si Option B retenue)")
    body(doc, "Rappel : l'interface SAGE n'est requise que si VERSUS BANK choisit de conserver SAGE comme "
              "moteur de paie (Option B, cf. ch. 2.4). En Option A (paie native NexusRH), ce chapitre est sans "
              "objet. Si Option B : decrire l'architecture du connecteur : flux SIRH->SAGE (donnees "
              "contractuelles, individuelles, variables de paie), flux SAGE->SIRH (bulletins PDF, journaux, "
              "statut de traitement). Preciser le mode d'echange (API REST, fichiers CSV/Excel, SFTP, connecteur "
              "natif), la frequence (batch/temps reel), la gestion des rejets et la reprise, ainsi que la "
              "separation des taches (SoD).")
    h3(doc, "4.2  Annuaire / IAM")
    body(doc, "Integration AD/Azure AD : SSO, provisioning, synchronisation des attributs, cycle de vie des comptes.")
    h3(doc, "4.3  Autres integrations")
    body(doc, "E-learning (inscriptions, completions, attestations) ; comptabilite/finance (masse salariale, "
              "couts RH, provisions) ; systeme bancaire (referentiel agences, controle interne). API documentees "
              "(Swagger), supervision, tracabilite des flux, environnements de test.")
    h3(doc, "4.4  Strategie de migration des donnees")
    body(doc, "Cartographie des sources (Excel, SAGE, papier numerise) ; regles de nettoyage, dedoublonnage, "
              "validation ; volumetrie et historique a reprendre par domaine ; protocole de recette des donnees "
              "migrees ; strategie de bascule et de rollback.")

    # ---- Ch.5 ----
    doc.add_page_break()
    h1(doc, "Chapitre 5 - Methodologie et planning (10 points)")
    body(doc, "Proposer un planning realiste sur 9 mois avec jalons et livrables, aligne sur les modalites de "
              "paiement du CCAP (20% cadrage, 40% VABF, 30% reception provisoire, 10% reception definitive).")
    table(doc, ["Phase / Lot", "Periode", "Livrables", "Jalon paiement"],
          [["Cadrage et conception", "M1", "SFD, DAT, dossier de parametrage", "20% (caution)"],
           ["Lot 1 - Socle + paie/SAGE + securite", "M1-M4", "Admin, dossier, absences, interface SAGE, classification/SSO/SIEM, organigramme, self-service", ""],
           ["Recette fonctionnelle Lot 1 (VABF)", "M4", "PV VABF, integration SAGE validee", "40%"],
           ["Lot 2 - Recrutement/talents/formation", "M4-M7", "Onboarding/offboarding, disciplinaire, talents, formation, signature", ""],
           ["Lot 3 - Avance + climat + reporting", "M7-M9", "Successions/mobilites/calibrage, climat social, reporting avance", ""],
           ["Formation + Go-Live (reception provisoire)", "M9", "Guides utilisateurs, mise en production", "30%"],
           ["Garantie / VSR (reception definitive)", "M9+12", "Levee des reserves", "10%"]],
          widths_cm=[4.6, 2.2, 7.3, 2.4])
    body(doc, "Approche : ateliers de cadrage, conception fonctionnelle/technique, parametrage/developpement "
              "specifique, migration, tests (unitaires, integration, recette), formation, mise en production, "
              "support au demarrage. Gouvernance : comites de pilotage/projet, MOA = DRH, MOE = OpenLab en "
              "coordination avec la DSI.")

    # ---- Ch.6 ----
    h1(doc, "Chapitre 6 - Conduite du changement, formation et support (10 points)")
    bullet(doc, "cartographie des processus, identification des irritants, co-construction avec les utilisateurs, communication (tutoriels, videos, FAQ), ambassadeurs internes ;", "Conduite du changement : ")
    bullet(doc, "plan par profil (RH, managers, collaborateurs, administrateurs), transfert de competences, guides utilisateurs et d'exploitation ;", "Formation : ")
    bullet(doc, "niveaux N1/N2/N3, prise en charge P1 < 2h et resolution < 8h ouvrees (exigence CCAP), canal mail/telephone/portail, support local WhatsApp OpenLab Abidjan ;", "Support et SLA : ")
    bullet(doc, "corrective, evolutive ; calendrier de releases, gestion des regressions, environnements de test.", "Maintenance : ")

    # ---- Ch.7 ----
    h1(doc, "Chapitre 7 - Equipe dediee et references (10 points)")
    body(doc, "[A RENSEIGNER] Fournir l'organigramme projet et les CV : chef de projet, expert fonctionnel "
              "RH/paie CI, architecte/lead technique, expert securite, responsable conduite du changement, "
              "support. Joindre au moins 3 references de SIRH integres (RH+Paie+Talents) des 5 dernieres "
              "annees, dont au moins une dans le secteur bancaire/financier. Demontrer l'assise financiere "
              "(CA 3 ans) et la perennite de l'editeur OpenLab Consulting.")

    # ---- Ch.8 ----
    doc.add_page_break()
    h1(doc, "Chapitre 8 - Reponses aux questions du DAO (1 a 10)")
    body(doc, "Le DAO pose 10 series de questions. Repondre point par point selon le format : (i) Oui/Non "
              "standard, (ii) description et limites, (iii) prerequis, (iv) elements de preuve, (v) impacts "
              "cout/delai/risque.")
    table(doc, ["Bloc de questions DAO", "Renvoi dans l'offre"],
          [["1. Perimetre et strategie de deploiement (big bang vs lots, roadmap 9 mois)", "Ch. 1.3 + Ch. 5"],
           ["2. Paie (SAGE) et gestion des variables (flux, modes d'echange, SoD)", "Ch. 2.4 + Ch. 4.1"],
           ["3. Habilitations, confidentialite, classification 4 niveaux", "Ch. 3.3"],
           ["4. Securite technique, audit et conformite (SSO, chiffrement, SIEM, PRA)", "Ch. 3.2 + 3.4"],
           ["5. Donnees, migration et qualite", "Ch. 4.4"],
           ["6. Integrations SI (AD/SSO, e-learning, finance, systeme bancaire)", "Ch. 4.2 + 4.3"],
           ["7. Points fonctionnels (onboarding videos, disciplinaire, signature, organigramme, Bloom)", "Ch. 2"],
           ["8. KPI, reporting et pilotage (definition, droits, adoption)", "Ch. 1.2 + 2.7"],
           ["9. Exploitation, support et SLA (N1/N2/N3, P1/P2/P3, supervision)", "Ch. 6"],
           ["10. Reversibilite, propriete des donnees et livrables", "Ch. 9"]],
          widths_cm=[12.5, 4.0])

    # ---- Ch.9 ----
    h1(doc, "Chapitre 9 - Reversibilite, propriete des donnees et livrables")
    bullet(doc, "VERSUS BANK reste proprietaire exclusif des donnees ; OpenLab accorde un droit d'utilisation (licence/abonnement) sur le logiciel ;", "Propriete : ")
    bullet(doc, "a la fin du marche, restitution sous 30 jours dans un format structure lisible (fichiers plats, SQL, PDF) + effacement securise des environnements, sans surcout de licence ;", "Reversibilite (CCAP) : ")
    bullet(doc, "SFD, DAT, dossier de parametrage, plans et rapports de tests, guides utilisateurs (RH/managers/collaborateurs), dossier d'exploitation, contrat de maintenance et SLA ;", "Livrables : ")
    bullet(doc, "formation des administrateurs SIRH (fonctionnel + technique), capacite a faire evoluer les workflows en interne.", "Transfert de competences : ")

    doc.add_paragraph()
    fin = doc.add_paragraph()
    run(fin, "Canevas genere pour OpenLab Consulting - trame de redaction de l'offre technique en reponse au "
             "DAO DRH/SIRH/29042026. Les sections [A RENSEIGNER] doivent etre completees avant soumission.",
        size=8.5, italic=True, color=GRAY)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    out = build()
    print("OK:", out, "-", os.path.getsize(out), "bytes")
